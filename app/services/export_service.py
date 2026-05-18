"""Reusable export helpers for generated research outputs."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
import re
from typing import Any, Iterable

from app.orchestrator.task_schema import TaskType


class ExportServiceError(RuntimeError):
    """Raised when an export file cannot be generated."""


@dataclass(frozen=True)
class ExportedFile:
    """Metadata for a generated export file."""

    format: str
    path: Path
    filename: str
    mime_type: str


@dataclass(frozen=True)
class MarkdownBlock:
    kind: str
    text: str | None = None
    level: int = 0
    rows: list[list[str]] | None = None


class ExportService:
    """Save research output as Markdown, DOCX, and PDF files."""

    MIME_TYPES = {
        "md": "text/markdown",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf": "application/pdf",
    }

    def __init__(self, output_dir: str | Path = "data/outputs") -> None:
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export_all(
        self,
        *,
        content: str,
        task_type: TaskType | str,
        title: str | None = None,
        created_on: date | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> dict[str, ExportedFile]:
        """Export content to all supported formats and return generated paths."""

        return {
            extension: self.export(
                content=content,
                task_type=task_type,
                extension=extension,
                title=title,
                created_on=created_on,
                metadata=metadata,
            )
            for extension in ("md", "docx", "pdf")
        }

    def export(
        self,
        *,
        content: str,
        task_type: TaskType | str,
        extension: str,
        title: str | None = None,
        created_on: date | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> ExportedFile:
        """Export content to one supported format."""

        extension = extension.lower().lstrip(".")
        if extension not in self.MIME_TYPES:
            raise ExportServiceError(f"Unsupported export format: {extension}")

        cleaned_content = content.strip()
        if not cleaned_content:
            raise ExportServiceError("Cannot export empty content.")

        normalized_task = self._normalize_task_type(task_type)
        export_title = (title or self._infer_title(cleaned_content, metadata) or "research_output").strip()
        filename = self.build_filename(
            title=export_title,
            task_type=normalized_task,
            extension=extension,
            created_on=created_on,
        )
        path = self.output_dir / filename

        try:
            if extension == "md":
                self._write_markdown(path=path, content=cleaned_content)
            elif extension == "docx":
                self._write_docx(path=path, content=cleaned_content, title=export_title)
            elif extension == "pdf":
                self._write_pdf(path=path, content=cleaned_content, title=export_title)
        except ImportError as exc:
            raise ExportServiceError(f"Missing dependency for {extension} export: {exc.name}") from exc

        return ExportedFile(
            format=extension,
            path=path,
            filename=filename,
            mime_type=self.MIME_TYPES[extension],
        )

    def build_filename(
        self,
        *,
        title: str,
        task_type: TaskType | str,
        extension: str,
        created_on: date | None = None,
    ) -> str:
        """Build a stable export filename from title, date, and task type."""

        normalized_task = self._normalize_task_type(task_type)
        output_date = created_on or date.today()
        safe_title = self._slugify(title) or "research-output"
        return f"{safe_title}_{output_date.isoformat()}_{normalized_task}.{extension.lstrip('.')}"

    def _write_markdown(self, *, path: Path, content: str) -> None:
        path.write_text(content + "\n", encoding="utf-8")

    def _write_docx(self, *, path: Path, content: str, title: str) -> None:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.shared import Inches, Pt

        document = Document()
        document.core_properties.title = title
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(10.5)

        section = document.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

        blocks = list(self._parse_markdown(content))
        if any(block.kind == "table" and block.rows and len(block.rows[0]) > 4 for block in blocks):
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

        for block in blocks:
            if block.kind == "heading":
                document.add_heading(block.text or "", level=min(max(block.level, 1), 4))
            elif block.kind == "bullet":
                document.add_paragraph(block.text or "", style="List Bullet")
            elif block.kind == "table" and block.rows:
                table = document.add_table(rows=0, cols=len(block.rows[0]))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"
                for row_index, row_cells in enumerate(block.rows):
                    cells = table.add_row().cells
                    for index, value in enumerate(row_cells):
                        cells[index].text = value
                        cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                        for paragraph in cells[index].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8.5 if len(row_cells) > 4 else 9.5)
                                run.bold = row_index == 0
                document.add_paragraph()
            elif block.text:
                document.add_paragraph(block.text)

        document.save(path)

    def _write_pdf(self, *, path: Path, content: str, title: str) -> None:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

        blocks = list(self._parse_markdown(content))
        wide_tables = [block for block in blocks if block.kind == "table" and block.rows and len(block.rows[0]) > 4]
        page_size = landscape(A4) if wide_tables else A4
        document = SimpleDocTemplate(
            str(path),
            pagesize=page_size,
            title=title,
            rightMargin=0.45 * inch,
            leftMargin=0.45 * inch,
            topMargin=0.5 * inch,
            bottomMargin=0.5 * inch,
        )
        styles = getSampleStyleSheet()
        body = ParagraphStyle(
            "ResearchBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            alignment=TA_LEFT,
            spaceAfter=6,
        )
        bullet = ParagraphStyle("ResearchBullet", parent=body, leftIndent=14, bulletIndent=4)
        heading_styles = {
            1: styles["Heading1"],
            2: styles["Heading2"],
            3: styles["Heading3"],
            4: styles["Heading4"],
        }
        story: list[Any] = []

        for block in blocks:
            if block.kind == "heading":
                story.append(Paragraph(self._pdf_escape(block.text or ""), heading_styles.get(block.level, styles["Heading4"])))
            elif block.kind == "bullet":
                story.append(Paragraph(self._pdf_escape(block.text or ""), bullet, bulletText="-"))
            elif block.kind == "table" and block.rows:
                col_count = len(block.rows[0])
                usable_width = page_size[0] - document.leftMargin - document.rightMargin
                col_widths = [usable_width / col_count] * col_count
                table_data = [
                    [Paragraph(self._pdf_escape(cell), body) for cell in row]
                    for row in block.rows
                ]
                table = Table(table_data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e9eef5")),
                            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#8a96a3")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 4),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]
                    )
                )
                story.extend([table, Spacer(1, 8)])
            elif block.text:
                story.append(Paragraph(self._pdf_escape(block.text), body))

        document.build(story)

    def _parse_markdown(self, content: str) -> Iterable[MarkdownBlock]:
        lines = content.splitlines()
        paragraph: list[str] = []
        index = 0

        def flush_paragraph() -> MarkdownBlock | None:
            if not paragraph:
                return None
            text = " ".join(part.strip() for part in paragraph if part.strip()).strip()
            paragraph.clear()
            return MarkdownBlock(kind="paragraph", text=text) if text else None

        while index < len(lines):
            line = lines[index].rstrip()
            stripped = line.strip()

            if not stripped:
                block = flush_paragraph()
                if block:
                    yield block
                index += 1
                continue

            table_rows = self._collect_table(lines, index)
            if table_rows:
                block = flush_paragraph()
                if block:
                    yield block
                yield MarkdownBlock(kind="table", rows=table_rows)
                index += self._table_line_count(lines, index)
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if heading_match:
                block = flush_paragraph()
                if block:
                    yield block
                yield MarkdownBlock(
                    kind="heading",
                    level=len(heading_match.group(1)),
                    text=self._strip_inline_markdown(heading_match.group(2)),
                )
                index += 1
                continue

            bullet_match = re.match(r"^[-*+]\s+(.+)$", stripped)
            if bullet_match:
                block = flush_paragraph()
                if block:
                    yield block
                yield MarkdownBlock(kind="bullet", text=self._strip_inline_markdown(bullet_match.group(1)))
                index += 1
                continue

            paragraph.append(self._strip_inline_markdown(stripped))
            index += 1

        block = flush_paragraph()
        if block:
            yield block

    def _collect_table(self, lines: list[str], start_index: int) -> list[list[str]] | None:
        if start_index + 1 >= len(lines):
            return None
        header = lines[start_index].strip()
        separator = lines[start_index + 1].strip()
        if not self._is_table_row(header) or not self._is_table_separator(separator):
            return None

        rows = [self._split_table_row(header)]
        index = start_index + 2
        while index < len(lines) and self._is_table_row(lines[index].strip()):
            row = self._split_table_row(lines[index].strip())
            if len(row) == len(rows[0]):
                rows.append(row)
            index += 1
        return rows

    def _table_line_count(self, lines: list[str], start_index: int) -> int:
        count = 2
        index = start_index + 2
        while index < len(lines) and self._is_table_row(lines[index].strip()):
            count += 1
            index += 1
        return count

    @staticmethod
    def _is_table_row(line: str) -> bool:
        return line.startswith("|") and line.endswith("|") and line.count("|") >= 2

    @staticmethod
    def _is_table_separator(line: str) -> bool:
        if not ExportService._is_table_row(line):
            return False
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)

    @staticmethod
    def _split_table_row(line: str) -> list[str]:
        return [ExportService._strip_inline_markdown(cell.strip()) for cell in line.strip("|").split("|")]

    @staticmethod
    def _strip_inline_markdown(text: str) -> str:
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"_([^_]+)_", r"\1", text)
        return text.strip()

    @staticmethod
    def _pdf_escape(text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

    @staticmethod
    def _normalize_task_type(task_type: TaskType | str) -> str:
        if isinstance(task_type, TaskType):
            return task_type.value
        return str(task_type).strip().lower().replace(" ", "_") or "research_output"

    @staticmethod
    def _infer_title(content: str, metadata: dict[str, Any] | None) -> str | None:
        if metadata:
            for key in ("title", "source_file"):
                value = metadata.get(key)
                if isinstance(value, str) and value.strip():
                    return Path(value).stem
            document = metadata.get("document")
            if isinstance(document, dict):
                source_file = document.get("source_file")
                if isinstance(source_file, str) and source_file.strip():
                    return Path(source_file).stem

        for line in content.splitlines():
            heading_match = re.match(r"^#{1,6}\s+(.+)$", line.strip())
            if heading_match:
                return ExportService._strip_inline_markdown(heading_match.group(1))
        return None

    @staticmethod
    def _slugify(value: str) -> str:
        slug = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower())
        return slug.strip("-")[:80]
